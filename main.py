import os
from math import floor
from typing import TypedDict

import imgkit
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Mm


class _PageSize(TypedDict):
    """
    Define os campos necessários para formatar o tamanho da página.
    """

    page_height: int
    page_width: int
    left_margin: int
    right_margin: int
    top_margin: int
    bottom_margin: int
    header_distance: int
    footer_distance: int


# A4: 210 x 297 mm
A4: _PageSize = {
    "page_height": 297,
    "page_width": 210,
    "left_margin": 25.4,
    "right_margin": 25.4,
    "top_margin": 25.4,
    "bottom_margin": 25.4,
    "header_distance": 12.7,
    "footer_distance": 12.7,
}

# Comando para abrir o arquivo gerado
OPEN_FILE_CMD = "open"


class DocumentSigner:
    """
    Classe responsável por adicionar um banner de assinatura em um documento Word.
    """

    def __init__(self, filename: str, page_size: _PageSize = A4) -> None:
        self.document = Document(filename)
        self.page_size = page_size

        for section in self.document.sections:
            for key, value in self.page_size.items():
                setattr(section, key, Mm(value))

    @staticmethod
    def __mm2px(mm: float, dpi: int = 96) -> int:
        """
        Converte milímetros para pixels.
        """

        return floor(mm * dpi / 25.4)

    def __get_usable_page_width(self) -> int:
        """
        Retorna a largura útil da página, ou seja, a largura da página menos as margens.
        """

        return self.page_size["page_width"] - (
            self.page_size["left_margin"] + self.page_size["right_margin"]
        )

    def __get_imgkit_config(self) -> dict[str, str]:
        """
        Retorna as configurações necessárias para o imgkit.
        """

        return {
            "format": "png",
            "enable-local-file-access": "",
            # 15 = padding arbitrário, para garantir que a imagem não fique cortada
            "crop-w": DocumentSigner.__mm2px(self.__get_usable_page_width()) + 15,
            # AVISO: A imagem pode ficar manchada com fundo transparente
            "transparent": "",
        }

    def __create_sign_banner(self, link: str, uuid: str) -> None:
        """
        Cria o banner a partir de um arquivo HTML. No caso concreto, o banner será
        criado a partir de from_url que receberá `link` e `uuid` como parâmetros.
        """

        imgkit.from_file(
            "./includes/sign_banner.html",
            "build/banner.png",
            options=self.__get_imgkit_config(),
        )

    def __populate_footer(self, link: str, uuid: str) -> None:
        """
        Preenche o footer do documento
        """

        self.__create_sign_banner(link, uuid)
        print(self.__get_usable_page_width())

        for section in self.document.sections:
            footer = section.footer

            # limpa todos os parágrafos do footer
            for paragraph in footer.paragraphs:
                paragraph.clear()

            paragraph = footer.paragraphs[0]
            p_run = paragraph.add_run()

            banner = p_run.add_picture(
                "build/banner.png",
                width=Mm(self.__get_usable_page_width()),
            )
            r_id = paragraph.part.relate_to(
                link,
                RELATIONSHIP_TYPE.HYPERLINK,
                is_external=True,
            )

            hyperlink = OxmlElement("a:hlinkClick")
            hyperlink.set(qn("r:id"), r_id)

            banner._inline.docPr.append(hyperlink)

    def sign(self, link: str, uuid: str) -> None:
        """
        Adiciona um banner de assinatura no footer do documento e salva uma cópia
        do documento modificado.
        """

        self.__populate_footer(link, uuid)
        self.document.save("build/file_out.docx")


if __name__ == "__main__":
    signer = DocumentSigner("assets/file_in.docx")

    signer.sign(
        link="https://solar.defensoria.to.def.br/docs/d/validar/",
        uuid="A6B56B39D2-195AD85977-740AB544E6-CB1D22BD03",
    )

    os.system(f"{OPEN_FILE_CMD} build/file_out.docx")
    exit(0)
